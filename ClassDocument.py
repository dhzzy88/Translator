from docx import Document

import BaiduTranslation
from docx.shared import Inches
from BaiduTranslation import returnback
import os.path
import string
import time

class myDocument(object):
    def __init__(self,path):
        document = Document(path)
        document_copy = Document()
        copy_dir = os.path.dirname(path)
        try:
            for oneParagraph in document.paragraphs:
               
 
                #将每个段落的格式和原来的格式设置一样，尽量保证翻译后的文档可以获得相同的阅读效果
                #tmparagraph = document_copy.add_paragraph(oneParagraph.text,style=oneParagraph.style)
                tmparagraph = document_copy.add_paragraph()
                tmparagraph.paragraph_format.alignment = oneParagraph.paragraph_format.alignment
                tmparagraph.paragraph_format.space_after = oneParagraph.paragraph_format.space_after
                tmparagraph.paragraph_format.left_indent = oneParagraph.paragraph_format.left_indent
                tmparagraph.paragraph_format.right_indent = oneParagraph.paragraph_format.right_indent
                tmparagraph.paragraph_format.line_spacing = oneParagraph.paragraph_format.line_spacing
                tmparagraph.paragraph_format.page_break_before = oneParagraph.paragraph_format.page_break_before
                
                #遍历段落中的runs,并且将其中的文本拿出来翻译
                for runn in oneParagraph.runs:
                    time.sleep(1)
                    run_result =runn
                   
                    #TODO(sun dhzzy88@163.com):translate 把翻译结果传递给result,将每个run中的文字翻译后传递到此处result
                    #run_result.text= filter(lambda x: x in string.printable, run_result.text)
                    if run_result.text in ["\n","\r","\t","\v"]:
                        run_result.text = runn.text
                    elif len(run_result.text)==0:
                        run_result.text = runn.text
                    else:
                        run_result.text = returnback(runn.text)
                
                    #print("df",len(run_result.text))
                    tmprunn = tmparagraph.add_run(run_result.text)
                   
                    # TODO(2019.10.18):设置每个run的格式，以便获得相同的显示结果
                    tmprunn.font.name = runn.font.name
                    tmprunn.font.size = runn.font.size
                    tmprunn.font.shadow = runn.font.shadow
                    tmprunn.font.strike = runn.font.strike
                    tmprunn.font.color.rgb = runn.font.color.rgb
                    tmprunn.font.bold = runn.font.bold
                    tmprunn.font.superscript = runn.font.superscript
                    tmprunn.font.subscript = runn.font.subscript
                   
                    print(tmprunn.text)

            #将获得的文件存储在原来的文件夹    
            document_copy.save( copy_dir + "\\new.docx" )
        except Exception as e:
            print(e)
        finally:
           #TODO(dhzzy88@163.com):计划增加一些对于文档处理结束后的操作
           pass




if __name__ == "__main__":
    path = "c:\\Users\\zhaozhiyi\\Desktop\\woo03.docx"
    doc = myDocument(path)